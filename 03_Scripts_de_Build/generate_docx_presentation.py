import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_docx():
    doc = docx.Document()
    
    # Page setup (Margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    # Primary Heading
    title = doc.add_heading('Apresentação Técnica: LeetCode #3', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    sub = doc.add_paragraph('Longest Substring Without Repeating Characters · Documento Editável da Apresentação')
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(4, 120, 87)
    sub.runs[0].font.bold = True

    meta = doc.add_paragraph('IFSP São Carlos · Tópicos em Algoritmos e Estruturas de Dados · Prof. Dr. Rodrigo Elias Bianchi\nIntegrantes: Rafael Feltrim, Ian, Gustavo (Gub)\n')
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph('—' * 55)

    slides_data = [
        {
            "num": "Slide 01",
            "title": "Capa Oficial",
            "content": [
                ("Título Principal", "Longest Substring Without Repeating Characters"),
                ("Subtítulo", "Desafio Técnico de Processos Seletivos (LeetCode #3) — Otimização Assintótica de O(N³) para O(N), Benchmark de CPU e Análise da Janela Deslizante."),
                ("Equipe", "Rafael Feltrim (Dev/Apresentação), Ian (Pesquisa/Algoritmos), Gustavo Gub (Complexidade/C)"),
                ("Disciplina", "Tópicos em Algoritmos e Estruturas de Dados — IFSP São Carlos 2026")
            ]
        },
        {
            "num": "Slide 02",
            "title": "Relevância de Mercado & Big Techs",
            "content": [
                ("Contexto", "Problema clássico nas entrevistas da Amazon (Top 10 SDE), Meta (Triagem 45m), Google (L3/L4), Microsoft, Apple e Bloomberg."),
                ("Regra dos 70/30", "Em entrevistas de Big Techs, o código representa apenas 30% da nota. Os outros 70% avaliam decomposição lógica, comunicação clara e justificativa matemática de Big-O."),
                ("Critérios Avaliados", "Identificação da ineficiência da força bruta, domínio de Hash Maps/Arrays diretos e tratamento de casos de borda.")
            ]
        },
        {
            "num": "Slide 03",
            "title": "Enunciado Formal & Conceitos",
            "content": [
                ("Enunciado", "Dada uma string s, encontre o comprimento da maior substring contígua que não contenha caracteres repetidos."),
                ("Distinção Crítica", "Substring: estritamente contínua na memória (ex: 'abc' em 'abcde'). Subsequência: mantém a ordem mas não é contínua (ex: 'ace')."),
                ("Casos de Teste", "Ex 1: 'abcabcbb' -> 3 ('abc') | Ex 2: 'bbbbb' -> 1 ('b') | Ex 3: 'pwwkew' -> 3 ('wke')")
            ]
        },
        {
            "num": "Slide 04",
            "title": "As 3 Abordagens em Código (Comparação)",
            "content": [
                ("1. Força Bruta O(N³)", "def brute_force(s):\n    n = len(s)\n    max_len = 0\n    for i in range(n):\n        for j in range(i+1, n+1):\n            sub = s[i:j]\n            if len(set(sub)) == len(sub):\n                max_len = max(max_len, len(sub))\n    return max_len\n[Status: Rejeitado / Causa TLE]"),
                ("2. Sliding Window com Set O(2N)", "def sliding_window_set(s):\n    char_set = set()\n    left = max_len = 0\n    for right in range(len(s)):\n        while s[right] in char_set:\n            char_set.remove(s[left])\n            left += 1\n        char_set.add(s[right])\n        max_len = max(max_len, right - left + 1)\n    return max_len\n[Status: Contratado / Remoção de 1 em 1]"),
                ("3. Sliding Window com Hash Map O(N)", "def sliding_window_map(s):\n    char_map = {}\n    left = max_len = 0\n    for right, char in enumerate(s):\n        if char in char_map and char_map[char] >= left:\n            left = char_map[char] + 1  # Salto direto O(1)\n        char_map[char] = right\n        max_len = max(max_len, right - left + 1)\n    return max_len\n[Status: Forte Candidato / Padrão Ouro]")
            ]
        },
        {
            "num": "Slide 05",
            "title": "Execução Passo a Passo na String 'abcabcbb'",
            "content": [
                ("Passo 1", "Char 'a' | R=0, L=0 | Janela 'a' | char_map={'a':0} | max=1"),
                ("Passo 2", "Char 'b' | R=1, L=0 | Janela 'ab' | char_map={'a':0, 'b':1} | max=2"),
                ("Passo 3", "Char 'c' | R=2, L=0 | Janela 'abc' | char_map={'a':0, 'b':1, 'c':2} | max=3 (Pico Máximo)"),
                ("Passo 4", "Char 'a' (duplicado) | R=3, L salta 0->1 | Janela 'bca' | char_map={'a':3, 'b':1, 'c':2} | max=3"),
                ("Passo 5", "Char 'b' (duplicado) | R=4, L salta 1->2 | Janela 'cab' | char_map={'a':3, 'b':4, 'c':2} | max=3"),
                ("Passo 6", "Char 'c' (duplicado) | R=5, L salta 2->3 | Janela 'abc' | char_map={'a':3, 'b':4, 'c':5} | max=3"),
                ("Passo 7", "Char 'b' (duplicado no idx 4) | R=6, L salta 3->5 | Janela 'cb' | char_map={'a':3, 'b':6, 'c':5} | max=3"),
                ("Passo 8", "Char 'b' (duplicado no idx 6) | R=7, L salta 5->7 | Janela 'b' | char_map={'a':3, 'b':7, 'c':5} | max=3")
            ]
        },
        {
            "num": "Slide 06",
            "title": "Benchmark Experimental de CPU (Economia de 4 Horas para 11ms)",
            "content": [
                ("Métrica Hero", "-99,99999% de tempo de CPU | 416.000.000x menos instruções para N=50.000"),
                ("N = 100", "Força Bruta: 4.85 ms (171.700 ops) | Sliding Window Map: 0.032 ms (100 ops) [-99.94%]"),
                ("N = 1.000", "Força Bruta: 1.953 ms (~2.0 s) | Sliding Window Map: 0.232 ms (1.000 ops) [-99.99%]"),
                ("N = 5.000", "Força Bruta: ~15.0 s (20.8 bi ops) | Sliding Window Map: 1.113 ms (5.000 ops) [-99.999%]"),
                ("N = 50.000", "Força Bruta: ~250 min (> 4 HORAS) | Sliding Window Map: 11.1 ms (50.000 ops) [> 99.99999%]")
            ]
        },
        {
            "num": "Slide 07",
            "title": "Arquitetura Visual da Janela Deslizante & Hash Map",
            "content": [
                ("Ponteiro Direito (R)", "Avança sequencialmente lendo novos caracteres e expandindo a janela."),
                ("Tabela Hash (char_map)", "Armazena a relação char -> último índice em O(1)."),
                ("Ponteiro Esquerdo (L)", "Ao detectar duplicata dentro da janela, salta instantaneamente para map[char] + 1 sem retrocesso."),
                ("Fórmula de Tamanho", "tamanho_atual = R - L + 1")
            ]
        },
        {
            "num": "Slide 08",
            "title": "Implementação em C com Aritmética de Ponteiros (Código do Gub)",
            "content": [
                ("Código Central em C", "int ultima_pos[256];\nfor (int k = 0; k < 256; k++) ultima_pos[k] = -1;\n\nchar *esquerda = s;\nint melhor_tam = 0;\n\nfor (char *direita = s; *direita != '\\0'; direita++) {\n    unsigned char c = (unsigned char)*direita;\n    int idx_dir = (int)(direita - s);\n\n    if (ultima_pos[c] != -1 && ultima_pos[c] >= (int)(esquerda - s)) {\n        esquerda = s + ultima_pos[c] + 1; // Pulo direto em O(1)\n    }\n\n    ultima_pos[c] = idx_dir;\n    int tam_atual = (int)(direita - esquerda) + 1;\n    if (tam_atual > melhor_tam) melhor_tam = tam_atual;\n}"),
                ("Engenharia de Memória", "1. Aritmética de ponteiros pura: (direita - s) sem custo extra.\n2. Espaço O(1) estrito: 256 * sizeof(int) = 1.024 bytes (1 KB) na Stack.\n3. Zero chamadas a malloc() e zero risco de memory leak.\n4. Memória contígua L1 Cache-friendly.")
            ]
        },
        {
            "num": "Slide 09",
            "title": "Tabela Rigorosa de Complexidade (Big-O)",
            "content": [
                ("Força Bruta", "Tempo: O(N³) | Espaço: O(min(N, Σ)) | Status: Rejeitado"),
                ("Sliding Window (Set)", "Tempo: O(2N) = O(N) | Espaço: O(min(N, Σ)) | Status: Contratado"),
                ("Sliding Window (Hash Map)", "Tempo: O(N) Estrito | Espaço: O(min(N, Σ)) | Status: Forte Candidato (Strong Hire)"),
                ("Justificativa de Tempo", "O ponteiro direito avança de 0 a N-1 apenas 1 vez. Consultas e inserções são O(1)."),
                ("Justificativa de Espaço", "A memória é limitada pelo menor valor entre o tamanho da string N e o alfabeto ASCII (Σ = 256).")
            ]
        },
        {
            "num": "Slide 10",
            "title": "Casos de Borda & Encerramento",
            "content": [
                ("Casos de Borda", "1. String Vazia \"\": retorna 0 imediatamente.\n2. Caracteres Idênticos \"bbbbbb\": janela mantém tamanho 1.\n3. Todos Distintos \"abcdef\": janela expande até o tamanho total N.\n4. Espaços e Símbolos \"a b c!\": tratados nativamente pela tabela ASCII de 256 bytes."),
                ("Conclusão", "Apresentação da Equipe (Rafael Feltrim, Ian, Gub) para a disciplina do Prof. Dr. Rodrigo Elias Bianchi. Repositório no GitHub e Vault Obsidian totalmente sincronizados.")
            ]
        }
    ]

    for slide in slides_data:
        h = doc.add_heading(f"{slide['num']} — {slide['title']}", level=1)
        h.runs[0].font.size = Pt(16)
        h.runs[0].font.color.rgb = RGBColor(9, 13, 22)
        h.runs[0].font.bold = True

        table = doc.add_table(rows=len(slide["content"]), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for i, (label, val) in enumerate(slide["content"]):
            row = table.rows[i]
            cell_label = row.cells[0]
            cell_val = row.cells[1]
            
            cell_label.width = Inches(1.8)
            cell_val.width = Inches(4.8)

            set_cell_background(cell_label, "F1F5F9")
            set_cell_background(cell_val, "FFFFFF")

            p_lbl = cell_label.paragraphs[0]
            p_lbl.text = label
            p_lbl.runs[0].font.bold = True
            p_lbl.runs[0].font.size = Pt(10)
            p_lbl.runs[0].font.color.rgb = RGBColor(51, 65, 85)

            p_val = cell_val.paragraphs[0]
            p_val.text = val
            p_val.runs[0].font.size = Pt(10)
            if "def " in val or "int " in val or "for " in val:
                p_val.runs[0].font.name = 'Consolas'
                p_val.runs[0].font.size = Pt(9)
                set_cell_background(cell_val, "F8FAFC")

        doc.add_paragraph('')

    output_docx = "Apresentacao_LeetCode3_Editavel.docx"
    doc.save(output_docx)
    print(f"Documento DOCX editável gerado com sucesso: {output_docx}")

if __name__ == "__main__":
    create_docx()
